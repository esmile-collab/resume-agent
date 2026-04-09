# Input: FastAPI 应用、测试客户端和 SQLite 重置环境。
# Output: 验证 HTTP API 端到端行为。
# Pos: 现行 API 验收测试。
# Rule: 一旦我被更新，务必同步更新本文件头注释与所属目录 README。
"""Acceptance tests for the session/message HTTP API."""

from __future__ import annotations

import os
import shutil
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from src.api.main import app
from src.db.database import DB_PATH, init_db

ARTIFACTS_ROOT = Path(".data") / "artifacts"
UPLOADS_ROOT = Path(".data") / "uploads"


def setup_function() -> None:
    """Reset DB and artifacts before each test."""
    if os.path.exists(DB_PATH):
        os.remove(DB_PATH)
    if ARTIFACTS_ROOT.exists():
        shutil.rmtree(ARTIFACTS_ROOT)
    if UPLOADS_ROOT.exists():
        shutil.rmtree(UPLOADS_ROOT)
    init_db().close()


def test_agent_root_endpoint() -> None:
    """Should expose a browser-friendly root endpoint."""
    client = TestClient(app)

    response = client.get("/")

    assert response.status_code == 200
    assert response.json()["service"] == "resume-agent"
    assert response.json()["health"] == "/health"
    assert response.json()["api_prefix"] == "/agent"


def test_agent_api_loop_end_to_end() -> None:
    """Should create a session, ingest JD, score, generate, and expose traces."""
    client = TestClient(app)

    started = client.post(
        "/agent/sessions",
        json={
            "project_name": "api-agent-e2e",
            "cycle": "2026秋招",
            "base_resume_text": "工作经历\n负责增长分析和项目推进\n教育背景\n本科",
        },
    )
    assert started.status_code == 200
    session_id = started.json()["session_id"]

    info = client.post(
        f"/agent/sessions/{session_id}/messages",
        json={
            "role": "user",
            "content": "我想投策略产品，我有2年产品经验，在某互联网公司负责增长项目，提升转化率20%。",
        },
    )
    assert info.status_code == 200
    assert "策略产品" in info.json()["reply"]

    ingest = client.post(
        f"/agent/sessions/{session_id}/messages",
        json={
            "role": "user",
            "content": "这是策略产品 JD",
            "attachments": [
                {
                    "type": "jd",
                    "name": "strategy_pm.txt",
                    "content": (
                        "岗位职责：负责增长策略制定和商业化分析。"
                        "任职要求：2年以上产品经验，数据分析能力强。"
                    ),
                }
            ],
        },
    )
    assert ingest.status_code == 200
    assert ingest.json()["intent"]["intent"] == "ingest_jd"

    score = client.post(
        f"/agent/sessions/{session_id}/messages",
        json={"role": "user", "content": "请帮我评分"},
    )
    assert score.status_code == 200
    assert score.json()["intent"]["intent"] == "score_resume"
    assert "score_report" in score.json()["tool_steps"][0]["observation"]

    generate = client.post(
        f"/agent/sessions/{session_id}/messages",
        json={"role": "user", "content": "请生成简历"},
    )
    assert generate.status_code == 200
    generated = generate.json()["tool_steps"][-1]["observation"]
    assert generated["ok"] is True
    assert Path(generated["artifact"]["path"]).exists()

    snapshot = client.get(f"/agent/sessions/{session_id}")
    assert snapshot.status_code == 200
    trace_kinds = [item["kind"] for item in snapshot.json()["traces"]]
    assert "intent" in trace_kinds
    assert "tool_call" in trace_kinds
    assert "observation" in trace_kinds


def test_agent_tools_endpoint() -> None:
    """Should expose the tool registry for frontend hydration."""
    client = TestClient(app)

    response = client.get("/agent/tools")

    assert response.status_code == 200
    tool_names = [item["name"] for item in response.json()["tools"]]
    assert "resume_generate" in tool_names
    assert "resume_polish" in tool_names


def test_agent_upload_endpoint() -> None:
    """Should accept multipart upload and return normalized attachment payload."""
    client = TestClient(app)
    started = client.post("/agent/sessions", json={"project_name": "upload-agent"})
    assert started.status_code == 200
    session_id = started.json()["session_id"]

    response = client.post(
        "/agent/uploads",
        data={"attachment_type": "jd", "session_id": session_id},
        files={"file": ("strategy_pm.txt", "岗位职责：负责增长策略。", "text/plain")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["type"] == "jd"
    assert payload["name"] == "strategy_pm.txt"
    assert "岗位职责" in payload["content"]
    assert payload["detected_type"] == "jd"
    assert payload["path"].endswith("strategy_pm.txt")
    assert Path(payload["path"]).is_absolute()


def test_agent_upload_endpoint_can_auto_detect_resume() -> None:
    """Should infer resume attachments when the user does not preselect the type."""
    client = TestClient(app)

    response = client.post(
        "/agent/uploads",
        data={"attachment_type": "other"},
        files={
            "file": (
                "candidate_profile.md",
                "个人总结\n有两年产品经验\n工作经历\n负责增长策略\n教育背景\n本科",
                "text/markdown",
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["type"] == "resume"
    assert payload["detected_type"] == "resume"
    assert "工作经历" in payload["content"]


def test_agent_upload_endpoint_can_parse_docx_resume() -> None:
    """Should extract text from docx resumes and infer the resume type."""
    client = TestClient(app)

    document = Document()
    document.add_heading("候选人简历", level=1)
    document.add_paragraph("个人总结")
    document.add_paragraph("有两年产品经验，负责增长策略与数据分析。")
    document.add_paragraph("工作经历")
    document.add_paragraph("负责用户转化优化，提升转化率20%。")
    document.add_paragraph("教育背景")
    document.add_paragraph("本科")
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = client.post(
        "/agent/uploads",
        data={"attachment_type": "other"},
        files={
            "file": (
                "candidate_resume.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["type"] == "resume"
    assert payload["detected_type"] == "resume"
    assert "工作经历" in payload["content"]
    assert "提升转化率20%" in payload["content"]


def test_agent_management_endpoints() -> None:
    """Should support manual track, JD, profile, and experience management."""
    client = TestClient(app)
    started = client.post("/agent/sessions", json={"project_name": "manage-agent"})
    assert started.status_code == 200
    session_payload = started.json()
    session_id = session_payload["session_id"]
    project_id = session_payload["project_id"]

    created_track = client.post(
        f"/agent/projects/{project_id}/tracks",
        json={
            "name": "商业分析",
            "positioning": "聚焦商业分析与增长策略岗位。",
            "core_keywords": ["商业分析", "增长", "SQL"],
            "resume_strategy": "优先突出分析和策略落地。",
            "default_resume_outline": "1. 摘要\n2. 经历\n3. 技能",
        },
    )
    assert created_track.status_code == 200
    track_id = created_track.json()["track"]["id"]

    jd_one = client.post(
        f"/agent/tracks/{track_id}/jds",
        json={
            "name": "ba_jd_1.txt",
            "content": "岗位职责：负责商业分析与增长策略。任职要求：熟悉 SQL。",
            "set_as_primary": True,
        },
    )
    assert jd_one.status_code == 200
    jd_one_id = jd_one.json()["jd"]["id"]
    assert jd_one.json()["jd"]["is_primary"] is True

    jd_two = client.post(
        f"/agent/tracks/{track_id}/jds",
        json={
            "name": "ba_jd_2.txt",
            "content": "岗位职责：负责经营分析和数据建模。任职要求：能独立产出洞察。",
            "set_as_primary": False,
        },
    )
    assert jd_two.status_code == 200
    jd_two_id = jd_two.json()["jd"]["id"]

    listed = client.get(f"/agent/tracks/{track_id}/jds")
    assert listed.status_code == 200
    assert len(listed.json()["jds"]) == 2

    switched = client.post(
        f"/agent/tracks/{track_id}/primary-jd",
        json={"jd_entry_id": jd_two_id},
    )
    assert switched.status_code == 200
    assert switched.json()["track"]["primary_jd_entry_id"] == jd_two_id

    renamed_jd = client.patch(
        f"/agent/jds/{jd_two_id}",
        json={"name": "ba_main_jd.txt", "content": "岗位职责：负责经营分析。任职要求：具备 SQL 能力。"},
    )
    assert renamed_jd.status_code == 200
    assert renamed_jd.json()["jd"]["name"] == "ba_main_jd.txt"

    updated_track = client.patch(
        f"/agent/tracks/{track_id}",
        json={
            "name": "经营分析",
            "positioning": "聚焦经营分析、增长和策略支持。",
            "core_keywords": ["经营分析", "增长", "SQL", "建模"],
            "resume_strategy": "突出分析能力和业务理解。",
            "default_resume_outline": "1. 摘要\n2. 重点项目\n3. 技能",
        },
    )
    assert updated_track.status_code == 200
    assert updated_track.json()["track"]["name"] == "经营分析"

    updated_profile = client.patch(
        f"/agent/projects/{project_id}/profile",
        json={
            "summary": "2 年商业分析经验，关注增长与经营决策。",
            "basics": {"education": "本科", "years_of_experience": "2"},
            "preferences": {"target_roles": ["经营分析"], "preferred_city": "上海"},
            "constraints": {"exclusions": ["纯销售岗位"]},
        },
    )
    assert updated_profile.status_code == 200
    assert updated_profile.json()["profile"]["summary"].startswith("2 年商业分析经验")

    created_experience = client.post(
        f"/agent/projects/{project_id}/experiences",
        json={
            "title": "经营分析实习",
            "organization": "某互联网公司",
            "time_range": "2025.01 - 2025.06",
            "summary": "负责渠道投放分析，优化预算分配。",
            "tags": ["经营分析", "投放"],
            "metrics": ["20%"],
            "evidence": ["周报复盘"],
            "confidence": 0.9,
            "source": "manual",
        },
    )
    assert created_experience.status_code == 200
    experience_id = created_experience.json()["experience"]["id"]

    updated_experience = client.patch(
        f"/agent/experiences/{experience_id}",
        json={
            "title": "经营分析项目",
            "organization": "某互联网公司",
            "time_range": "2025.01 - 2025.06",
            "summary": "负责渠道投放分析并推动预算优化。",
            "tags": ["经营分析", "预算"],
            "metrics": ["20%"],
            "evidence": ["周报复盘", "项目总结"],
            "confidence": 0.95,
            "source": "manual",
        },
    )
    assert updated_experience.status_code == 200
    assert updated_experience.json()["experience"]["title"] == "经营分析项目"

    snapshot = client.get(f"/agent/sessions/{session_id}")
    assert snapshot.status_code == 200
    payload = snapshot.json()["snapshot"]
    assert payload["profile"]["basics"]["education"] == "本科"
    assert payload["tracks"][0]["name"] == "经营分析"
    assert payload["tracks"][0]["primary_jd_entry_id"] == jd_two_id
    assert payload["experiences"][0]["title"] == "经营分析项目"

    deleted_jd = client.delete(f"/agent/jds/{jd_one_id}")
    assert deleted_jd.status_code == 200

    deleted_experience = client.delete(f"/agent/experiences/{experience_id}")
    assert deleted_experience.status_code == 200

    deleted_track = client.delete(f"/agent/tracks/{track_id}")
    assert deleted_track.status_code == 200

    final_snapshot = client.get(f"/agent/sessions/{session_id}")
    assert final_snapshot.status_code == 200
    assert final_snapshot.json()["snapshot"]["tracks"] == []
    assert final_snapshot.json()["snapshot"]["experiences"] == []


def test_agent_session_history_and_artifact_endpoints() -> None:
    """Should expose session history, artifact preview/diff, revision save, and export."""
    client = TestClient(app)

    started = client.post(
        "/agent/sessions",
        json={
            "project_name": "artifact-agent",
            "cycle": "2026秋招",
            "base_resume_text": "工作经历\n负责数据分析和增长策略\n教育背景\n本科",
        },
    )
    assert started.status_code == 200
    payload = started.json()
    session_id = payload["session_id"]
    project_id = payload["project_id"]

    client.post(
        f"/agent/sessions/{session_id}/messages",
        json={"role": "user", "content": "我想投商业分析，有2年分析经验。"},
    )
    ingest = client.post(
        f"/agent/sessions/{session_id}/messages",
        json={
            "role": "user",
            "content": "这是商业分析 JD",
            "attachments": [
                {
                    "type": "jd",
                    "name": "ba_jd.txt",
                    "content": "岗位职责：负责经营分析。任职要求：SQL、数据分析、策略拆解。",
                }
            ],
        },
    )
    assert ingest.status_code == 200
    track_id = ingest.json()["tool_steps"][0]["observation"]["track_id"]

    first_score = client.post(
        f"/agent/sessions/{session_id}/messages",
        json={"role": "user", "content": "请帮我评分", "active_track_id": track_id},
    )
    assert first_score.status_code == 200
    score_artifact_id = first_score.json()["tool_steps"][0]["observation"]["artifact"]["id"]

    resume_update = client.post(
        f"/agent/sessions/{session_id}/messages",
        json={
            "role": "user",
            "content": "我补充一版简历",
            "attachments": [
                {
                    "type": "resume",
                    "name": "resume_v2.txt",
                    "content": "工作经历\n负责经营分析、SQL 建模和预算优化，提升转化率15%\n教育背景\n本科",
                }
            ],
        },
    )
    assert resume_update.status_code == 200

    second_score = client.post(
        f"/agent/sessions/{session_id}/messages",
        json={"role": "user", "content": "再评分一次", "active_track_id": track_id},
    )
    assert second_score.status_code == 200
    latest_score_artifact_id = second_score.json()["tool_steps"][0]["observation"]["artifact"]["id"]
    assert latest_score_artifact_id != score_artifact_id

    generate = client.post(
        f"/agent/sessions/{session_id}/messages",
        json={"role": "user", "content": "请生成简历", "active_track_id": track_id},
    )
    assert generate.status_code == 200
    generated_artifact_id = generate.json()["tool_steps"][-1]["observation"]["artifact"]["id"]

    session_history = client.get("/agent/sessions")
    assert session_history.status_code == 200
    assert any(item["id"] == session_id for item in session_history.json()["sessions"])

    artifacts = client.get(f"/agent/projects/{project_id}/artifacts")
    assert artifacts.status_code == 200
    artifact_types = [item["artifact_type"] for item in artifacts.json()["artifacts"]]
    assert "score_report" in artifact_types
    assert "generated_resume" in artifact_types

    artifact_detail = client.get(f"/agent/artifacts/{latest_score_artifact_id}")
    assert artifact_detail.status_code == 200
    assert artifact_detail.json()["parsed_payload"]["final_score"] >= 0
    assert "简历评分报告" in artifact_detail.json()["content"]

    artifact_diff = client.post(
        f"/agent/artifacts/{latest_score_artifact_id}/diff",
        json={"base_artifact_id": score_artifact_id},
    )
    assert artifact_diff.status_code == 200
    assert "stats" in artifact_diff.json()

    export_docx = client.post(
        f"/agent/artifacts/{generated_artifact_id}/export",
        json={"format": "docx"},
    )
    assert export_docx.status_code == 200
    assert Path(export_docx.json()["path"]).exists()

    export_pdf = client.post(
        f"/agent/artifacts/{generated_artifact_id}/export",
        json={"format": "pdf"},
    )
    assert export_pdf.status_code == 200
    assert Path(export_pdf.json()["path"]).exists()

    revised = client.post(
        f"/agent/artifacts/{generated_artifact_id}/revisions",
        json={
            "content": "# 商业分析 定制简历\n\n## 相关经历\n- 手动补充一条更聚焦经营分析的 bullet。",
        },
    )
    assert revised.status_code == 200
    revised_artifact = revised.json()["artifact"]
    assert revised_artifact["artifact_type"] == "edited_resume"
    assert Path(revised_artifact["path"]).exists()
