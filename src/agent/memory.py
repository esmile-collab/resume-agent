# Input: SQLite CRUD、对话压缩器、项目目录结构和文件系统。
# Output: 输出项目快照、结构化记忆读写和产物管理能力。
# Pos: Agent 的长期记忆与资产中枢。
# Rule: 一旦我被更新，务必同步更新本文件头注释与所属目录 README。
"""Memory manager for profile, experiences, job tracks, and artifacts."""

from __future__ import annotations

import json
import re
from datetime import datetime, UTC
from difflib import unified_diff
from pathlib import Path
from typing import Any

from db.agent_crud import (
    AgentArtifactCRUD,
    AgentSessionCRUD,
    CandidateProfileCRUD,
    ExperienceItemCRUD,
    JobTrackCRUD,
    JobTrackJDLinkCRUD,
    SessionMessageCRUD,
)
from db.crud import ProjectCRUD, ProjectJDEntryCRUD
from docx import Document
from observability import DialogCompressionManager
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas


def _loads_json(text: str, *, default: Any) -> Any:
    try:
        return json.loads(text)
    except (TypeError, json.JSONDecodeError):
        return default


def _unique_keep_order(values: list[str]) -> list[str]:
    seen: set[str] = set()
    output: list[str] = []
    for value in values:
        normalized = value.strip()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        output.append(normalized)
    return output


_pdf_font_registered = False


class MemoryManager:
    """Load and update the structured memory layers used by the runtime."""

    ROLE_HINTS = {
        "策略产品": ["策略产品", "增长产品", "商业分析", "strategy product", "growth"],
        "功能产品": ["功能产品", "产品经理", "prd", "需求分析", "feature product"],
        "内容运营": ["内容运营", "新媒体", "社区运营", "用户运营", "内容策划"],
        "数据分析": ["数据分析", "商业分析", "data analyst", "bi", "数据产品"],
    }

    def __init__(self) -> None:
        self.project_crud = ProjectCRUD()
        self.project_jd_crud = ProjectJDEntryCRUD()
        self.session_crud = AgentSessionCRUD()
        self.profile_crud = CandidateProfileCRUD()
        self.experience_crud = ExperienceItemCRUD()
        self.track_crud = JobTrackCRUD()
        self.track_jd_link_crud = JobTrackJDLinkCRUD()
        self.message_crud = SessionMessageCRUD()
        self.trace_artifact_crud = AgentArtifactCRUD()
        self.dialog_manager = DialogCompressionManager()

    def build_snapshot(self, *, project_id: str, session_id: str = "") -> dict[str, Any]:
        profile_row = self.profile_crud.get(project_id)
        profile = {
            "summary": profile_row.summary if profile_row else "",
            "basics": _loads_json(profile_row.basics_json, default={}) if profile_row else {},
            "preferences": _loads_json(profile_row.preferences_json, default={}) if profile_row else {},
            "constraints": _loads_json(profile_row.constraints_json, default={}) if profile_row else {},
        }

        experiences = []
        for item in self.experience_crud.list_by_project(project_id):
            experiences.append(
                {
                    "id": item.id,
                    "title": item.title,
                    "organization": item.organization,
                    "time_range": item.time_range,
                    "summary": item.summary,
                    "tags": _loads_json(item.tags_json, default=[]),
                    "metrics": _loads_json(item.metrics_json, default=[]),
                    "evidence": _loads_json(item.evidence_json, default=[]),
                    "confidence": item.confidence,
                    "source": item.source,
                }
            )

        tracks = []
        for track in self.track_crud.list_by_project(project_id):
            links = self.track_jd_link_crud.list_by_track(track.id)
            tracks.append(
                {
                    "id": track.id,
                    "name": track.name,
                    "positioning": track.positioning,
                    "core_keywords": _loads_json(track.core_keywords_json, default=[]),
                    "resume_strategy": track.resume_strategy,
                    "default_resume_outline": track.default_resume_outline,
                    "primary_jd_entry_id": track.primary_jd_entry_id,
                    "jd_count": len(links),
                    "jd_ids": [link.project_jd_entry_id for link in links],
                }
            )

        artifacts = []
        for artifact in self.trace_artifact_crud.list_by_project(project_id, limit=10):
            artifacts.append(
                {
                    "id": artifact.id,
                    "track_id": artifact.track_id,
                    "jd_entry_id": artifact.jd_entry_id,
                    "artifact_type": artifact.artifact_type,
                    "version": artifact.version,
                    "path": artifact.path,
                    "summary": _loads_json(artifact.summary_json, default={}),
                    "created_at": artifact.created_at.isoformat() if artifact.created_at else "",
                }
            )

        recent_messages: list[dict[str, Any]] = []
        if session_id:
            for message in self.message_crud.list_by_session(session_id, limit=20):
                recent_messages.append(
                    {
                        "id": message.id,
                        "role": message.role,
                        "content": message.content,
                        "metadata": _loads_json(message.metadata_json, default={}),
                    }
                )

        summary = {}
        summary_path = self._project_root(project_id) / "state" / "dialog_summary.json"
        if summary_path.exists():
            summary = self.dialog_manager.read_summary(project_root=self._project_root(project_id))

        return {
            "profile": profile,
            "experiences": experiences,
            "tracks": tracks,
            "artifacts": artifacts,
            "recent_messages": recent_messages,
            "dialog_summary": summary,
        }

    def ensure_project_dirs(self, project_id: str) -> None:
        """Create the project-local directories required by the new runtime."""
        root = self._project_root(project_id)
        for rel in ("jd", "resume", "state", "agent/outputs"):
            (root / rel).mkdir(parents=True, exist_ok=True)

    def apply_dialog_memory(
        self,
        *,
        project_id: str,
        content: str,
        attachments: list[dict[str, str]],
    ) -> dict[str, Any]:
        snapshot = self.build_snapshot(project_id=project_id)
        current_profile = snapshot["profile"]
        basics = dict(current_profile["basics"])
        preferences = dict(current_profile["preferences"])
        constraints = dict(current_profile["constraints"])
        summary_parts = [part for part in [current_profile["summary"]] if part]

        target_roles = _unique_keep_order(
            list(preferences.get("target_roles", [])) + self._extract_target_roles(content)
        )
        if target_roles:
            preferences["target_roles"] = target_roles
            summary_parts.append(f"目标方向：{' / '.join(target_roles)}")

        years_match = re.search(r"(\d+(?:\.\d+)?)\s*年(?:工作|经验|产品|运营|实习)?经验", content)
        if years_match:
            basics["years_of_experience"] = years_match.group(1)

        education_match = re.search(r"(本科|硕士|博士|大专|MBA)", content)
        if education_match:
            basics["education"] = education_match.group(1)

        city_match = re.search(r"(?:base在|在|希望去|目标城市|工作地)\s*[:：]?\s*([A-Za-z\u4e00-\u9fff]{2,12})", content)
        if city_match:
            preferences["preferred_city"] = city_match.group(1)

        if "不考虑" in content:
            constraints["exclusions"] = _unique_keep_order(
                list(constraints.get("exclusions", [])) + [content.strip()]
            )

        profile_updated = any(
            [
                target_roles,
                years_match is not None,
                education_match is not None,
                city_match is not None,
                "不考虑" in content,
            ]
        )
        if profile_updated:
            self.profile_crud.upsert(
                project_id,
                summary="；".join(_unique_keep_order(summary_parts)),
                basics=basics,
                preferences=preferences,
                constraints=constraints,
            )

        new_experience_ids: list[str] = []
        existing_summaries = {item["summary"] for item in snapshot["experiences"]}
        for extracted in self._extract_experiences(content):
            if extracted["summary"] in existing_summaries:
                continue
            item = self.experience_crud.create(
                project_id,
                title=extracted["title"],
                organization=extracted["organization"],
                time_range=extracted["time_range"],
                summary=extracted["summary"],
                tags=extracted["tags"],
                metrics=extracted["metrics"],
                evidence=[extracted["summary"]],
            )
            new_experience_ids.append(item.id)

        track_ids: list[str] = []
        for role in target_roles:
            track = self.ensure_track(project_id=project_id, track_name=role)
            track_ids.append(track["id"])

        resume_attachments = [item for item in attachments if item.get("type") == "resume"]
        if resume_attachments:
            self.write_resume_text(project_id, resume_attachments[-1].get("content", ""))

        return {
            "profile_updated": profile_updated,
            "new_experience_ids": new_experience_ids,
            "track_ids": track_ids,
        }

    def ensure_track(
        self,
        *,
        project_id: str,
        track_name: str,
        keywords: list[str] | None = None,
        positioning: str = "",
    ) -> dict[str, Any]:
        default_outline = (
            "1. 目标岗位摘要\n"
            "2. 最相关经历（2-3条）\n"
            "3. 关键技能与结果\n"
            "4. 教育背景/补充信息"
        )
        track = self.track_crud.get_or_create(
            project_id,
            name=track_name,
            positioning=positioning or f"面向 {track_name} 方向的长期求职资产",
            core_keywords=keywords or self._default_keywords_for_track(track_name),
            resume_strategy=f"优先突出与 {track_name} 直接相关的经历与可量化结果。",
            default_resume_outline=default_outline,
        )
        return {
            "id": track.id,
            "name": track.name,
            "positioning": track.positioning,
            "core_keywords": _loads_json(track.core_keywords_json, default=[]),
            "resume_strategy": track.resume_strategy,
            "default_resume_outline": track.default_resume_outline,
            "primary_jd_entry_id": track.primary_jd_entry_id,
        }

    def attach_jd_to_track(
        self,
        *,
        project_id: str,
        jd_text: str,
        source_name: str,
        track_name: str = "",
    ) -> dict[str, Any]:
        resolved_track_name = track_name or self.infer_track_name(jd_text)
        track = self.ensure_track(
            project_id=project_id,
            track_name=resolved_track_name,
            keywords=self.extract_keywords(jd_text)[:8],
        )
        jd_entry = self.project_jd_crud.create(project_id=project_id, content=jd_text, source_file=source_name)
        self.track_jd_link_crud.link(track["id"], jd_entry.id, metadata={"source_name": source_name})
        track_record = self.track_crud.get(track["id"])
        if track_record is not None and not track_record.primary_jd_entry_id:
            self.track_crud.update(
                track["id"],
                name=track_record.name,
                positioning=track_record.positioning,
                core_keywords=_loads_json(track_record.core_keywords_json, default=[]),
                resume_strategy=track_record.resume_strategy,
                default_resume_outline=track_record.default_resume_outline,
                primary_jd_entry_id=jd_entry.id,
            )

        jd_path = self._project_root(project_id) / "jd" / f"{jd_entry.id}.txt"
        jd_path.parent.mkdir(parents=True, exist_ok=True)
        jd_path.write_text(jd_text, encoding="utf-8")

        return {
            "track_id": track["id"],
            "track_name": track["name"],
            "jd_entry_id": jd_entry.id,
            "keywords": self.extract_keywords(jd_text)[:8],
            "source_name": source_name,
        }

    def latest_jd_for_track(self, track_id: str) -> dict[str, Any] | None:
        track = self.track_crud.get(track_id)
        if track is not None and track.primary_jd_entry_id:
            jd_entry = self.project_jd_crud.get(track.primary_jd_entry_id)
            if jd_entry is not None:
                return {
                    "id": jd_entry.id,
                    "project_id": jd_entry.project_id,
                    "raw_content": jd_entry.raw_content,
                    "source_file": jd_entry.source_file or "",
                }
        links = self.track_jd_link_crud.list_by_track(track_id)
        if not links:
            return None
        latest_link = links[-1]
        jd_entry = self.project_jd_crud.get(latest_link.project_jd_entry_id)
        if jd_entry is None:
            return None
        return {
            "id": jd_entry.id,
            "project_id": jd_entry.project_id,
            "raw_content": jd_entry.raw_content,
            "source_file": jd_entry.source_file or "",
        }

    def resolve_track(self, *, project_id: str, track_id: str = "", track_name: str = "") -> dict[str, Any] | None:
        snapshot = self.build_snapshot(project_id=project_id)
        tracks = snapshot["tracks"]
        if track_id:
            for track in tracks:
                if track["id"] == track_id:
                    return track
        if track_name:
            for track in tracks:
                if track["name"] == track_name:
                    return track
        if len(tracks) == 1:
            return tracks[0]
        return None

    def get_resume_text(self, project_id: str) -> str:
        resume_path = self._project_root(project_id) / "resume" / "base_resume_v1.txt"
        if resume_path.exists():
            return resume_path.read_text(encoding="utf-8")
        fallback_path = self._project_root(project_id) / "resume" / "base_resume_from_session.txt"
        if fallback_path.exists():
            return fallback_path.read_text(encoding="utf-8")
        return ""

    def choose_relevant_experiences(
        self,
        *,
        project_id: str,
        track: dict[str, Any],
        limit: int = 3,
    ) -> list[dict[str, Any]]:
        snapshot = self.build_snapshot(project_id=project_id)
        track_keywords = set(track.get("core_keywords", []))
        scored: list[tuple[int, dict[str, Any]]] = []
        for item in snapshot["experiences"]:
            tags = set(item.get("tags", []))
            content = f"{item.get('title', '')} {item.get('summary', '')}"
            overlap = len(track_keywords & tags)
            overlap += len(track_keywords & set(self.extract_keywords(content)))
            scored.append((overlap, item))
        scored.sort(key=lambda pair: pair[0], reverse=True)
        chosen = [item for _, item in scored[:limit] if item]
        if chosen:
            return chosen
        return snapshot["experiences"][:limit]

    def next_artifact_version(self, *, track_id: str, artifact_type: str) -> int:
        latest = self.trace_artifact_crud.latest_for_track(track_id, artifact_type)
        if latest is None:
            return 1
        return latest.version + 1

    def register_artifact(
        self,
        *,
        project_id: str,
        track_id: str,
        jd_entry_id: str,
        artifact_type: str,
        version: int,
        path: str,
        summary: dict[str, Any],
    ) -> dict[str, Any]:
        record = self.trace_artifact_crud.create(
            project_id=project_id,
            track_id=track_id,
            jd_entry_id=jd_entry_id,
            artifact_type=artifact_type,
            version=version,
            path=path,
            summary=summary,
        )
        return {
            "id": record.id,
            "track_id": record.track_id,
            "jd_entry_id": record.jd_entry_id,
            "artifact_type": record.artifact_type,
            "version": record.version,
            "path": record.path,
            "summary": summary,
            "created_at": record.created_at.isoformat() if record.created_at else "",
        }

    def append_project_dialog(
        self,
        *,
        project_id: str,
        role: str,
        content: str,
        facts: list[str] | None = None,
    ) -> dict[str, Any]:
        project_root = self._project_root(project_id)
        project_root.mkdir(parents=True, exist_ok=True)
        return self.dialog_manager.append_turn(
            project_root=project_root,
            role=role,
            content=content,
            facts=facts,
        )

    def infer_track_name(self, text: str) -> str:
        lowered = text.lower()
        for name, hints in self.ROLE_HINTS.items():
            if any(hint.lower() in lowered for hint in hints):
                return name
        if any(token in text for token in ("增长", "商业化", "策略")):
            return "策略产品"
        if any(token in text for token in ("需求", "功能", "PRD", "prd")):
            return "功能产品"
        return "通用方向"

    def extract_keywords(self, text: str) -> list[str]:
        tokens = re.findall(r"[A-Za-z][A-Za-z0-9+#.\-/]{2,}|[\u4e00-\u9fff]{2,8}", text.lower())
        stopwords = {"岗位", "职责", "要求", "负责", "具备", "经验", "能力", "进行", "相关"}
        output: list[str] = []
        for token in tokens:
            if token in stopwords or token in output:
                continue
            output.append(token)
        return output[:20]

    def _extract_target_roles(self, content: str) -> list[str]:
        roles: list[str] = []
        explicit = re.search(
            r"(?:求职方向|目标岗位|想投|目标是|我想投|我想找)\s*[:：]?\s*([^\n。；;]+)",
            content,
        )
        if explicit:
            explicit_text = explicit.group(1).strip()
            matched_known_role = False
            for role_name, hints in self.ROLE_HINTS.items():
                if any(hint.lower() in explicit_text.lower() for hint in hints):
                    roles.append(role_name)
                    matched_known_role = True
            if not matched_known_role:
                roles.extend(re.split(r"[、,/，和]+", explicit_text))
        lowered = content.lower()
        for role_name, hints in self.ROLE_HINTS.items():
            if any(hint.lower() in lowered for hint in hints):
                roles.append(role_name)
        cleaned = []
        for role in roles:
            normalized = role.strip("：:，,。. ")
            if len(normalized) < 2:
                continue
            if len(normalized) > 12 and normalized not in self.ROLE_HINTS:
                continue
            cleaned.append(normalized)
        return _unique_keep_order(cleaned)

    def _extract_experiences(self, content: str) -> list[dict[str, Any]]:
        chunks = [item.strip() for item in re.split(r"[\n]+", content) if item.strip()]
        experiences: list[dict[str, Any]] = []
        for chunk in chunks:
            if not any(token in chunk for token in ("项目", "实习", "工作", "负责", "主导", "推动")):
                continue
            title = chunk[:18].strip("：:，, ")
            metrics = re.findall(r"\d+%?|\d+\.\d+", chunk)
            tags = self.extract_keywords(chunk)[:6]
            time_range_match = re.search(r"(\d{4}[./-]\d{1,2}\s*[~至-]\s*\d{4}[./-]\d{1,2})", chunk)
            organization_match = re.search(r"在([\u4e00-\u9fffA-Za-z0-9]+)(?:实习|工作|项目)", chunk)
            experiences.append(
                {
                    "title": title if title else "对话补充经历",
                    "organization": organization_match.group(1) if organization_match else "",
                    "time_range": time_range_match.group(1) if time_range_match else "",
                    "summary": chunk,
                    "tags": tags,
                    "metrics": metrics,
                }
            )
        return experiences

    def _default_keywords_for_track(self, track_name: str) -> list[str]:
        return list(self.ROLE_HINTS.get(track_name, [track_name]))[:6]

    def _project_root(self, project_id: str) -> Path:
        return Path(".data") / "artifacts" / project_id

    def write_resume_text(self, project_id: str, resume_text: str) -> None:
        self.ensure_project_dirs(project_id)
        resume_path = self._project_root(project_id) / "resume" / "base_resume_from_session.txt"
        resume_path.parent.mkdir(parents=True, exist_ok=True)
        resume_path.write_text(resume_text, encoding="utf-8")

    def create_track(
        self,
        *,
        project_id: str,
        name: str,
        positioning: str,
        core_keywords: list[str],
        resume_strategy: str,
        default_resume_outline: str,
    ) -> dict[str, Any]:
        if not name.strip():
            raise ValueError("Track name is required.")
        existing = self.track_crud.get_by_name(project_id, name)
        if existing is not None:
            raise ValueError(f"Track already exists: {name}")
        track = self.track_crud.create(
            project_id,
            name=name,
            positioning=positioning,
            core_keywords=core_keywords,
            resume_strategy=resume_strategy,
            default_resume_outline=default_resume_outline,
        )
        return self._track_to_dict(track)

    def update_track(
        self,
        *,
        track_id: str,
        name: str,
        positioning: str,
        core_keywords: list[str],
        resume_strategy: str,
        default_resume_outline: str,
    ) -> dict[str, Any]:
        if not name.strip():
            raise ValueError("Track name is required.")
        current = self.track_crud.get(track_id)
        if current is None:
            raise ValueError(f"Track not found: {track_id}")
        duplicate = self.track_crud.get_by_name(current.project_id, name)
        if duplicate is not None and duplicate.id != track_id:
            raise ValueError(f"Track already exists: {name}")
        updated = self.track_crud.update(
            track_id,
            name=name,
            positioning=positioning,
            core_keywords=core_keywords,
            resume_strategy=resume_strategy,
            default_resume_outline=default_resume_outline,
            primary_jd_entry_id=current.primary_jd_entry_id,
        )
        if updated is None:
            raise ValueError(f"Track not found: {track_id}")
        return self._track_to_dict(updated)

    def delete_track(self, track_id: str) -> None:
        track = self.track_crud.get(track_id)
        if track is None:
            raise ValueError(f"Track not found: {track_id}")
        for link in list(self.track_jd_link_crud.list_by_track(track_id)):
            self.delete_jd(link.project_jd_entry_id)
        self.track_crud.delete(track_id)

    def list_track_jds(self, track_id: str) -> list[dict[str, Any]]:
        track = self.track_crud.get(track_id)
        if track is None:
            raise ValueError(f"Track not found: {track_id}")
        entries: list[dict[str, Any]] = []
        for link in self.track_jd_link_crud.list_by_track(track_id):
            jd_entry = self.project_jd_crud.get(link.project_jd_entry_id)
            if jd_entry is None:
                continue
            entries.append(
                {
                    "id": jd_entry.id,
                    "track_id": track_id,
                    "name": jd_entry.source_file or f"{jd_entry.id}.txt",
                    "content": jd_entry.raw_content,
                    "preview": jd_entry.raw_content[:180],
                    "created_at": jd_entry.created_at.isoformat() if jd_entry.created_at else "",
                    "is_primary": jd_entry.id == track.primary_jd_entry_id,
                }
            )
        return entries

    def create_jd(
        self,
        *,
        track_id: str,
        name: str,
        content: str,
        set_as_primary: bool = False,
    ) -> dict[str, Any]:
        track = self.track_crud.get(track_id)
        if track is None:
            raise ValueError(f"Track not found: {track_id}")
        if not content.strip():
            raise ValueError("JD content is required.")
        jd_entry = self.project_jd_crud.create(track.project_id, content=content, source_file=name.strip() or "jd.txt")
        self.track_jd_link_crud.link(track_id, jd_entry.id, metadata={"source_name": jd_entry.source_file})
        jd_path = self._project_root(track.project_id) / "jd" / f"{jd_entry.id}.txt"
        jd_path.parent.mkdir(parents=True, exist_ok=True)
        jd_path.write_text(content, encoding="utf-8")
        if set_as_primary or not track.primary_jd_entry_id:
            self.set_primary_jd(track_id=track_id, jd_entry_id=jd_entry.id)
        return self.list_track_jds(track_id)[-1]

    def update_jd(self, *, jd_entry_id: str, name: str, content: str | None = None) -> dict[str, Any]:
        if not name.strip():
            raise ValueError("JD name is required.")
        jd_entry = self.project_jd_crud.update(
            jd_entry_id,
            content=content,
            source_file=name.strip() or None,
        )
        if jd_entry is None:
            raise ValueError(f"JD not found: {jd_entry_id}")
        jd_path = self._project_root(jd_entry.project_id) / "jd" / f"{jd_entry.id}.txt"
        if content is not None:
            jd_path.parent.mkdir(parents=True, exist_ok=True)
            jd_path.write_text(content, encoding="utf-8")
        links = self.track_jd_link_crud.list_by_jd_entry(jd_entry_id)
        track_id = links[0].track_id if links else ""
        is_primary = False
        if track_id:
            track = self.track_crud.get(track_id)
            is_primary = track is not None and track.primary_jd_entry_id == jd_entry.id
        return {
            "id": jd_entry.id,
            "track_id": track_id,
            "name": jd_entry.source_file or f"{jd_entry.id}.txt",
            "content": jd_entry.raw_content,
            "preview": jd_entry.raw_content[:180],
            "created_at": jd_entry.created_at.isoformat() if jd_entry.created_at else "",
            "is_primary": is_primary,
        }

    def set_primary_jd(self, *, track_id: str, jd_entry_id: str) -> dict[str, Any]:
        track = self.track_crud.get(track_id)
        if track is None:
            raise ValueError(f"Track not found: {track_id}")
        valid_ids = {link.project_jd_entry_id for link in self.track_jd_link_crud.list_by_track(track_id)}
        if jd_entry_id not in valid_ids:
            raise ValueError(f"JD {jd_entry_id} is not linked to track {track_id}")
        updated = self.track_crud.update(
            track_id,
            name=track.name,
            positioning=track.positioning,
            core_keywords=_loads_json(track.core_keywords_json, default=[]),
            resume_strategy=track.resume_strategy,
            default_resume_outline=track.default_resume_outline,
            primary_jd_entry_id=jd_entry_id,
        )
        if updated is None:
            raise ValueError(f"Track not found: {track_id}")
        return self._track_to_dict(updated)

    def delete_jd(self, jd_entry_id: str) -> None:
        jd_entry = self.project_jd_crud.get(jd_entry_id)
        if jd_entry is None:
            raise ValueError(f"JD not found: {jd_entry_id}")
        links = self.track_jd_link_crud.list_by_jd_entry(jd_entry_id)
        for link in links:
            track = self.track_crud.get(link.track_id)
            self.track_jd_link_crud.delete(link.track_id, jd_entry_id)
            if track is None or track.primary_jd_entry_id != jd_entry_id:
                continue
            remaining_links = self.track_jd_link_crud.list_by_track(link.track_id)
            next_primary = remaining_links[-1].project_jd_entry_id if remaining_links else ""
            self.track_crud.update(
                track.id,
                name=track.name,
                positioning=track.positioning,
                core_keywords=_loads_json(track.core_keywords_json, default=[]),
                resume_strategy=track.resume_strategy,
                default_resume_outline=track.default_resume_outline,
                primary_jd_entry_id=next_primary,
            )
        self.project_jd_crud.delete(jd_entry_id)
        jd_path = self._project_root(jd_entry.project_id) / "jd" / f"{jd_entry_id}.txt"
        if jd_path.exists():
            jd_path.unlink()

    def update_profile(
        self,
        *,
        project_id: str,
        summary: str,
        basics: dict[str, Any],
        preferences: dict[str, Any],
        constraints: dict[str, Any],
    ) -> dict[str, Any]:
        profile = self.profile_crud.upsert(
            project_id,
            summary=summary,
            basics=basics,
            preferences=preferences,
            constraints=constraints,
        )
        return {
            "summary": profile.summary,
            "basics": _loads_json(profile.basics_json, default={}),
            "preferences": _loads_json(profile.preferences_json, default={}),
            "constraints": _loads_json(profile.constraints_json, default={}),
        }

    def create_experience(
        self,
        *,
        project_id: str,
        title: str,
        organization: str,
        time_range: str,
        summary: str,
        tags: list[str],
        metrics: list[str],
        evidence: list[str],
        confidence: float,
        source: str,
    ) -> dict[str, Any]:
        if not title.strip():
            raise ValueError("Experience title is required.")
        record = self.experience_crud.create(
            project_id,
            title=title,
            organization=organization,
            time_range=time_range,
            summary=summary,
            tags=tags,
            metrics=metrics,
            evidence=evidence,
            confidence=confidence,
            source=source,
        )
        return self._experience_to_dict(record)

    def update_experience(
        self,
        *,
        experience_id: str,
        title: str,
        organization: str,
        time_range: str,
        summary: str,
        tags: list[str],
        metrics: list[str],
        evidence: list[str],
        confidence: float,
        source: str,
    ) -> dict[str, Any]:
        if not title.strip():
            raise ValueError("Experience title is required.")
        record = self.experience_crud.update(
            experience_id,
            title=title,
            organization=organization,
            time_range=time_range,
            summary=summary,
            tags=tags,
            metrics=metrics,
            evidence=evidence,
            confidence=confidence,
            source=source,
        )
        if record is None:
            raise ValueError(f"Experience not found: {experience_id}")
        return self._experience_to_dict(record)

    def delete_experience(self, experience_id: str) -> None:
        deleted = self.experience_crud.delete(experience_id)
        if not deleted:
            raise ValueError(f"Experience not found: {experience_id}")

    def list_sessions(self, *, limit: int = 30) -> list[dict[str, Any]]:
        sessions: list[dict[str, Any]] = []
        for session in self.session_crud.list_recent(limit=limit):
            project = self.project_crud.get(session.project_id)
            recent_messages = self.message_crud.list_by_session(session.id, limit=1)
            preview = recent_messages[0].content if recent_messages else ""
            sessions.append(
                {
                    "id": session.id,
                    "project_id": session.project_id,
                    "project_name": project.name if project else session.project_id,
                    "title": session.title,
                    "status": session.status,
                    "preview": preview[:120],
                    "created_at": session.created_at.isoformat() if session.created_at else "",
                    "updated_at": session.updated_at.isoformat() if session.updated_at else "",
                }
            )
        return sessions

    def list_artifacts(self, *, project_id: str, limit: int = 50) -> list[dict[str, Any]]:
        output: list[dict[str, Any]] = []
        for artifact in self.trace_artifact_crud.list_by_project(project_id, limit=limit):
            output.append(
                {
                    "id": artifact.id,
                    "track_id": artifact.track_id,
                    "jd_entry_id": artifact.jd_entry_id,
                    "artifact_type": artifact.artifact_type,
                    "version": artifact.version,
                    "path": artifact.path,
                    "summary": _loads_json(artifact.summary_json, default={}),
                    "created_at": artifact.created_at.isoformat() if artifact.created_at else "",
                }
            )
        return output

    def get_artifact_detail(self, artifact_id: str) -> dict[str, Any]:
        artifact = self.trace_artifact_crud.get(artifact_id)
        if artifact is None:
            raise ValueError(f"Artifact not found: {artifact_id}")
        summary = _loads_json(artifact.summary_json, default={})
        path = Path(artifact.path)
        content = path.read_text(encoding="utf-8") if path.exists() else ""
        parsed_payload: dict[str, Any] | None = None
        report_json_path = str(summary.get("report_json_path", "")).strip()
        if report_json_path:
            json_path = Path(report_json_path)
            if json_path.exists():
                parsed_payload = _loads_json(json_path.read_text(encoding="utf-8"), default=None)
        return {
            "artifact": {
                "id": artifact.id,
                "track_id": artifact.track_id,
                "jd_entry_id": artifact.jd_entry_id,
                "artifact_type": artifact.artifact_type,
                "version": artifact.version,
                "path": artifact.path,
                "summary": summary,
                "created_at": artifact.created_at.isoformat() if artifact.created_at else "",
            },
            "content": content,
            "parsed_payload": parsed_payload,
        }

    def diff_artifacts(self, *, artifact_id: str, base_artifact_id: str) -> dict[str, Any]:
        target = self.get_artifact_detail(artifact_id)
        base = self.get_artifact_detail(base_artifact_id)
        diff_lines = list(
            unified_diff(
                base["content"].splitlines(),
                target["content"].splitlines(),
                fromfile=str(base["artifact"]["path"]),
                tofile=str(target["artifact"]["path"]),
                lineterm="",
            )
        )
        additions = sum(1 for line in diff_lines if line.startswith("+") and not line.startswith("+++"))
        deletions = sum(1 for line in diff_lines if line.startswith("-") and not line.startswith("---"))
        return {
            "artifact_id": artifact_id,
            "base_artifact_id": base_artifact_id,
            "diff": "\n".join(diff_lines),
            "stats": {
                "additions": additions,
                "deletions": deletions,
            },
        }

    def save_artifact_revision(self, *, artifact_id: str, content: str) -> dict[str, Any]:
        if not content.strip():
            raise ValueError("Edited artifact content is required.")

        artifact = self.trace_artifact_crud.get(artifact_id)
        if artifact is None:
            raise ValueError(f"Artifact not found: {artifact_id}")
        if artifact.artifact_type not in {"generated_resume", "polished_resume", "edited_resume"}:
            raise ValueError("Only resume artifacts can be edited manually.")

        artifact_type = "edited_resume"
        version = self.next_artifact_version(track_id=artifact.track_id, artifact_type=artifact_type)
        output_path = self._project_root(artifact.project_id) / "agent" / "outputs" / (
            f"{artifact.track_id}_edited_v{version}.md"
        )
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(content, encoding="utf-8")

        summary = _loads_json(artifact.summary_json, default={})
        summary.update(
            {
                "source_artifact_id": artifact.id,
                "char_count": len(content),
                "line_count": len([line for line in content.splitlines() if line.strip()]),
            }
        )
        return self.register_artifact(
            project_id=artifact.project_id,
            track_id=artifact.track_id,
            jd_entry_id=artifact.jd_entry_id,
            artifact_type=artifact_type,
            version=version,
            path=str(output_path.resolve()),
            summary=summary,
        )

    def export_artifact(self, *, artifact_id: str, export_format: str) -> dict[str, Any]:
        artifact = self.trace_artifact_crud.get(artifact_id)
        if artifact is None:
            raise ValueError(f"Artifact not found: {artifact_id}")
        source_path = Path(artifact.path)
        if not source_path.exists():
            raise ValueError(f"Artifact file not found: {artifact.path}")
        export_dir = source_path.parent / "exports"
        export_dir.mkdir(parents=True, exist_ok=True)
        export_path = export_dir / f"{source_path.stem}.{export_format}"
        content = source_path.read_text(encoding="utf-8")
        if export_format == "docx":
            self._write_docx(export_path=export_path, content=content)
        elif export_format == "pdf":
            self._write_pdf(export_path=export_path, content=content)
        else:
            raise ValueError(f"Unsupported export format: {export_format}")
        return {
            "artifact_id": artifact_id,
            "format": export_format,
            "path": str(export_path.resolve()),
            "generated_at": datetime.now(UTC).isoformat(),
        }

    def _experience_to_dict(self, item: Any) -> dict[str, Any]:
        return {
            "id": item.id,
            "title": item.title,
            "organization": item.organization,
            "time_range": item.time_range,
            "summary": item.summary,
            "tags": _loads_json(item.tags_json, default=[]),
            "metrics": _loads_json(item.metrics_json, default=[]),
            "evidence": _loads_json(item.evidence_json, default=[]),
            "confidence": item.confidence,
            "source": item.source,
        }

    def _track_to_dict(self, track: Any) -> dict[str, Any]:
        links = self.track_jd_link_crud.list_by_track(track.id)
        return {
            "id": track.id,
            "name": track.name,
            "positioning": track.positioning,
            "core_keywords": _loads_json(track.core_keywords_json, default=[]),
            "resume_strategy": track.resume_strategy,
            "default_resume_outline": track.default_resume_outline,
            "primary_jd_entry_id": track.primary_jd_entry_id,
            "jd_count": len(links),
            "jd_ids": [link.project_jd_entry_id for link in links],
        }

    @staticmethod
    def _write_docx(*, export_path: Path, content: str) -> None:
        document = Document()
        for raw_line in content.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            if line.startswith("# "):
                document.add_heading(line.removeprefix("# ").strip(), level=1)
            elif line.startswith("## "):
                document.add_heading(line.removeprefix("## ").strip(), level=2)
            elif line.startswith("### "):
                document.add_heading(line.removeprefix("### ").strip(), level=3)
            elif line.startswith("- "):
                document.add_paragraph(line.removeprefix("- ").strip(), style="List Bullet")
            else:
                document.add_paragraph(line)
        document.save(export_path)

    @staticmethod
    def _write_pdf(*, export_path: Path, content: str) -> None:
        global _pdf_font_registered
        if not _pdf_font_registered:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            _pdf_font_registered = True
        pdf = canvas.Canvas(str(export_path), pagesize=A4)
        width, height = A4
        y = height - 48
        pdf.setFont("STSong-Light", 12)
        for raw_line in content.splitlines():
            line = raw_line.strip()
            if not line:
                y -= 10
                continue
            for chunk in MemoryManager._wrap_text(line, max_chars=44):
                if y < 48:
                    pdf.showPage()
                    pdf.setFont("STSong-Light", 12)
                    y = height - 48
                pdf.drawString(48, y, chunk)
                y -= 18
        pdf.save()

    @staticmethod
    def _wrap_text(text: str, *, max_chars: int) -> list[str]:
        if len(text) <= max_chars:
            return [text]
        return [text[index:index + max_chars] for index in range(0, len(text), max_chars)]
