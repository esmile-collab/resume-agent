# Input: FastAPI、runtime、上传解析器与 pydantic 请求模型。
# Output: 输出 session/message API、管理端点和上传端点。
# Pos: 后端 HTTP 入口主文件。
# Rule: 一旦我被更新，务必同步更新本文件头注释与所属目录 README。
"""HTTP API for the session/message based resume agent runtime."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any, Literal

from agent import ResumeAgentRuntime
from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field, model_validator
from pypdf import PdfReader


class AgentAttachmentPayload(BaseModel):
    """Attachment payload for the new runtime API."""

    type: Literal["jd", "resume", "other"] = "other"
    content: str
    name: str = ""


class AgentSessionStartRequest(BaseModel):
    """Create or resume one runtime session."""

    project_id: str = ""
    project_name: str = ""
    cycle: str = ""
    base_resume_text: str = ""


class AgentMessageRequest(BaseModel):
    """Append one message into the new runtime session."""

    role: str = "user"
    content: str = ""
    attachments: list[AgentAttachmentPayload] = Field(default_factory=list)
    active_track_id: str = ""
    active_track_name: str = ""

    @model_validator(mode="after")
    def validate_input(self) -> "AgentMessageRequest":
        if not self.content.strip() and not self.attachments:
            raise ValueError("content or attachments is required")
        return self


class TrackUpsertRequest(BaseModel):
    name: str
    positioning: str = ""
    core_keywords: list[str] = Field(default_factory=list)
    resume_strategy: str = ""
    default_resume_outline: str = ""


class TrackPrimaryJDRequest(BaseModel):
    jd_entry_id: str


class JDCreateRequest(BaseModel):
    name: str = ""
    content: str
    set_as_primary: bool = False


class JDUpdateRequest(BaseModel):
    name: str = ""
    content: str | None = None


class CandidateProfileUpdateRequest(BaseModel):
    summary: str = ""
    basics: dict[str, Any] = Field(default_factory=dict)
    preferences: dict[str, Any] = Field(default_factory=dict)
    constraints: dict[str, Any] = Field(default_factory=dict)


class ExperienceUpsertRequest(BaseModel):
    title: str
    organization: str = ""
    time_range: str = ""
    summary: str = ""
    tags: list[str] = Field(default_factory=list)
    metrics: list[str] = Field(default_factory=list)
    evidence: list[str] = Field(default_factory=list)
    confidence: float = 0.7
    source: str = "manual"


class ArtifactDiffRequest(BaseModel):
    base_artifact_id: str


class ArtifactExportRequest(BaseModel):
    format: Literal["docx", "pdf"]


class ArtifactRevisionRequest(BaseModel):
    content: str


app = FastAPI(
    title="Resume Agent Local API",
    version="0.2.0",
    description="Session/message based resume agent API for the React frontend.",
)
app.add_middleware(
    CORSMiddleware,
    allow_origin_regex=r"^https?://(localhost|127\.0\.0\.1)(:\d+)?$",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
agent_runtime = ResumeAgentRuntime()


def _status_code_from_error(message: str) -> int:
    """Map common domain errors to HTTP status."""
    lowered = message.lower()
    if "not found" in lowered:
        return 404
    return 400


@app.get("/")
def root() -> dict[str, Any]:
    """Human-friendly root endpoint for local browser checks."""
    return {
        "service": "resume-agent",
        "status": "ok",
        "health": "/health",
        "docs": "/docs",
        "api_prefix": "/agent",
    }


@app.get("/health")
def health() -> dict[str, str]:
    """Basic liveness check for local runs."""
    return {"status": "ok"}


@app.post("/agent/sessions")
def create_agent_session(payload: AgentSessionStartRequest) -> dict[str, Any]:
    """Start a session for the new think-call-observe runtime."""
    try:
        return agent_runtime.start_session(
            project_id=payload.project_id,
            project_name=payload.project_name,
            cycle=payload.cycle,
            base_resume_text=payload.base_resume_text,
        )
    except ValueError as exc:
        raise HTTPException(status_code=_status_code_from_error(str(exc)), detail=str(exc)) from exc


@app.get("/agent/sessions")
def list_agent_sessions(limit: int = 30) -> dict[str, Any]:
    return agent_runtime.list_sessions(limit=limit)


@app.post("/agent/sessions/{session_id}/messages")
def create_agent_message(session_id: str, payload: AgentMessageRequest) -> dict[str, Any]:
    """Append one message and execute the server-side agent loop."""
    try:
        return agent_runtime.handle_message(
            session_id=session_id,
            role=payload.role,
            content=payload.content,
            attachments=[item.model_dump() for item in payload.attachments],
            active_track_id=payload.active_track_id,
            active_track_name=payload.active_track_name,
        )
    except ValueError as exc:
        raise HTTPException(status_code=_status_code_from_error(str(exc)), detail=str(exc)) from exc


@app.get("/agent/sessions/{session_id}")
def get_agent_session(session_id: str) -> dict[str, Any]:
    """Return one session snapshot for React hydration."""
    try:
        return agent_runtime.get_session_snapshot(session_id)
    except ValueError as exc:
        raise HTTPException(status_code=_status_code_from_error(str(exc)), detail=str(exc)) from exc


@app.get("/agent/tools")
def list_agent_tools() -> dict[str, Any]:
    """Expose tool catalog for the frontend."""
    return {"tools": agent_runtime.tools.catalog()}


@app.post("/agent/projects/{project_id}/tracks")
def create_track(project_id: str, payload: TrackUpsertRequest) -> dict[str, Any]:
    try:
        track = agent_runtime.memory.create_track(
            project_id=project_id,
            name=payload.name.strip(),
            positioning=payload.positioning.strip(),
            core_keywords=payload.core_keywords,
            resume_strategy=payload.resume_strategy.strip(),
            default_resume_outline=payload.default_resume_outline.strip(),
        )
        return {"track": track}
    except ValueError as exc:
        raise HTTPException(status_code=_status_code_from_error(str(exc)), detail=str(exc)) from exc


@app.patch("/agent/tracks/{track_id}")
def update_track(track_id: str, payload: TrackUpsertRequest) -> dict[str, Any]:
    try:
        track = agent_runtime.memory.update_track(
            track_id=track_id,
            name=payload.name.strip(),
            positioning=payload.positioning.strip(),
            core_keywords=payload.core_keywords,
            resume_strategy=payload.resume_strategy.strip(),
            default_resume_outline=payload.default_resume_outline.strip(),
        )
        return {"track": track}
    except ValueError as exc:
        raise HTTPException(status_code=_status_code_from_error(str(exc)), detail=str(exc)) from exc


@app.delete("/agent/tracks/{track_id}")
def delete_track(track_id: str) -> dict[str, Any]:
    try:
        agent_runtime.memory.delete_track(track_id)
        return {"ok": True}
    except ValueError as exc:
        raise HTTPException(status_code=_status_code_from_error(str(exc)), detail=str(exc)) from exc


@app.get("/agent/tracks/{track_id}/jds")
def list_track_jds(track_id: str) -> dict[str, Any]:
    try:
        return {"jds": agent_runtime.memory.list_track_jds(track_id)}
    except ValueError as exc:
        raise HTTPException(status_code=_status_code_from_error(str(exc)), detail=str(exc)) from exc


@app.post("/agent/tracks/{track_id}/jds")
def create_track_jd(track_id: str, payload: JDCreateRequest) -> dict[str, Any]:
    try:
        jd = agent_runtime.memory.create_jd(
            track_id=track_id,
            name=payload.name.strip() or "jd.txt",
            content=payload.content,
            set_as_primary=payload.set_as_primary,
        )
        return {"jd": jd}
    except ValueError as exc:
        raise HTTPException(status_code=_status_code_from_error(str(exc)), detail=str(exc)) from exc


@app.patch("/agent/jds/{jd_entry_id}")
def update_jd(jd_entry_id: str, payload: JDUpdateRequest) -> dict[str, Any]:
    try:
        jd = agent_runtime.memory.update_jd(
            jd_entry_id=jd_entry_id,
            name=payload.name.strip() or "jd.txt",
            content=payload.content,
        )
        return {"jd": jd}
    except ValueError as exc:
        raise HTTPException(status_code=_status_code_from_error(str(exc)), detail=str(exc)) from exc


@app.delete("/agent/jds/{jd_entry_id}")
def delete_jd(jd_entry_id: str) -> dict[str, Any]:
    try:
        agent_runtime.memory.delete_jd(jd_entry_id)
        return {"ok": True}
    except ValueError as exc:
        raise HTTPException(status_code=_status_code_from_error(str(exc)), detail=str(exc)) from exc


@app.post("/agent/tracks/{track_id}/primary-jd")
def set_track_primary_jd(track_id: str, payload: TrackPrimaryJDRequest) -> dict[str, Any]:
    try:
        track = agent_runtime.memory.set_primary_jd(track_id=track_id, jd_entry_id=payload.jd_entry_id)
        return {"track": track}
    except ValueError as exc:
        raise HTTPException(status_code=_status_code_from_error(str(exc)), detail=str(exc)) from exc


@app.patch("/agent/projects/{project_id}/profile")
def update_candidate_profile(project_id: str, payload: CandidateProfileUpdateRequest) -> dict[str, Any]:
    profile = agent_runtime.memory.update_profile(
        project_id=project_id,
        summary=payload.summary,
        basics=payload.basics,
        preferences=payload.preferences,
        constraints=payload.constraints,
    )
    return {"profile": profile}


@app.post("/agent/projects/{project_id}/experiences")
def create_experience(project_id: str, payload: ExperienceUpsertRequest) -> dict[str, Any]:
    try:
        experience = agent_runtime.memory.create_experience(
            project_id=project_id,
            title=payload.title.strip(),
            organization=payload.organization.strip(),
            time_range=payload.time_range.strip(),
            summary=payload.summary.strip(),
            tags=payload.tags,
            metrics=payload.metrics,
            evidence=payload.evidence,
            confidence=payload.confidence,
            source=payload.source.strip() or "manual",
        )
        return {"experience": experience}
    except ValueError as exc:
        raise HTTPException(status_code=_status_code_from_error(str(exc)), detail=str(exc)) from exc


@app.patch("/agent/experiences/{experience_id}")
def update_experience(experience_id: str, payload: ExperienceUpsertRequest) -> dict[str, Any]:
    try:
        experience = agent_runtime.memory.update_experience(
            experience_id=experience_id,
            title=payload.title.strip(),
            organization=payload.organization.strip(),
            time_range=payload.time_range.strip(),
            summary=payload.summary.strip(),
            tags=payload.tags,
            metrics=payload.metrics,
            evidence=payload.evidence,
            confidence=payload.confidence,
            source=payload.source.strip() or "manual",
        )
        return {"experience": experience}
    except ValueError as exc:
        raise HTTPException(status_code=_status_code_from_error(str(exc)), detail=str(exc)) from exc


@app.delete("/agent/experiences/{experience_id}")
def delete_experience(experience_id: str) -> dict[str, Any]:
    try:
        agent_runtime.memory.delete_experience(experience_id)
        return {"ok": True}
    except ValueError as exc:
        raise HTTPException(status_code=_status_code_from_error(str(exc)), detail=str(exc)) from exc


@app.get("/agent/projects/{project_id}/artifacts")
def list_project_artifacts(project_id: str, limit: int = 50) -> dict[str, Any]:
    return {"artifacts": agent_runtime.memory.list_artifacts(project_id=project_id, limit=limit)}


@app.get("/agent/artifacts/{artifact_id}")
def get_artifact_detail(artifact_id: str) -> dict[str, Any]:
    try:
        return agent_runtime.memory.get_artifact_detail(artifact_id)
    except ValueError as exc:
        raise HTTPException(status_code=_status_code_from_error(str(exc)), detail=str(exc)) from exc


@app.post("/agent/artifacts/{artifact_id}/diff")
def diff_artifact(artifact_id: str, payload: ArtifactDiffRequest) -> dict[str, Any]:
    try:
        return agent_runtime.memory.diff_artifacts(
            artifact_id=artifact_id,
            base_artifact_id=payload.base_artifact_id,
        )
    except ValueError as exc:
        raise HTTPException(status_code=_status_code_from_error(str(exc)), detail=str(exc)) from exc


@app.post("/agent/artifacts/{artifact_id}/export")
def export_artifact(artifact_id: str, payload: ArtifactExportRequest) -> dict[str, Any]:
    try:
        return agent_runtime.memory.export_artifact(artifact_id=artifact_id, export_format=payload.format)
    except ValueError as exc:
        raise HTTPException(status_code=_status_code_from_error(str(exc)), detail=str(exc)) from exc


@app.post("/agent/artifacts/{artifact_id}/revisions")
def save_artifact_revision(artifact_id: str, payload: ArtifactRevisionRequest) -> dict[str, Any]:
    try:
        artifact = agent_runtime.memory.save_artifact_revision(
            artifact_id=artifact_id,
            content=payload.content,
        )
        return {"artifact": artifact}
    except ValueError as exc:
        raise HTTPException(status_code=_status_code_from_error(str(exc)), detail=str(exc)) from exc


@app.post("/agent/uploads")
async def upload_agent_file(
    file: UploadFile = File(...),
    attachment_type: str = Form("other"),
    session_id: str = Form(""),
) -> dict[str, Any]:
    """Upload a real file and return one attachment payload for the chat composer."""
    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    safe_type = attachment_type if attachment_type in {"jd", "resume", "other"} else "other"
    text_content = _extract_upload_text(
        raw=raw,
        filename=file.filename or "upload.txt",
        content_type=file.content_type or "application/octet-stream",
    )
    detected_type = _infer_attachment_type(
        text_content=text_content,
        filename=file.filename or "upload.txt",
        content_type=file.content_type or "application/octet-stream",
    )
    effective_type = safe_type if safe_type != "other" else detected_type
    target_path = _persist_upload(raw=raw, filename=file.filename or "upload.txt", session_id=session_id)

    return {
        "type": effective_type,
        "detected_type": detected_type,
        "name": file.filename or "upload.txt",
        "content": text_content,
        "content_type": file.content_type or "application/octet-stream",
        "size": len(raw),
        "path": str(target_path),
    }


def _decode_upload_bytes(raw: bytes) -> str:
    """Decode plain-text bytes into text content."""
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _extract_upload_text(*, raw: bytes, filename: str, content_type: str) -> str:
    """Extract usable text from one uploaded file."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx" or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            document = Document(BytesIO(raw))
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            table_lines = []
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        table_lines.append(" | ".join(cells))
            text = "\n".join([*paragraphs, *table_lines]).strip()
            if text:
                return text
        except Exception:
            pass
        raise HTTPException(status_code=400, detail="DOCX 解析失败，当前文件无法提取文本。")
    if suffix == ".pdf" or content_type == "application/pdf":
        try:
            reader = PdfReader(BytesIO(raw))
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(part.strip() for part in pages if part.strip())
            if text.strip():
                return text
        except Exception:
            pass
        raise HTTPException(status_code=400, detail="PDF 解析失败，当前文件无法提取文本。")
    return _decode_upload_bytes(raw)


def _infer_attachment_type(*, text_content: str, filename: str, content_type: str) -> str:
    """Infer whether the uploaded file is a JD or a resume."""
    lowered_name = filename.lower()
    if any(token in lowered_name for token in ("resume", "cv", "简历")):
        return "resume"
    if any(token in lowered_name for token in ("jd", "job", "岗位", "职位")):
        return "jd"
    if content_type == "application/pdf" and "resume" in lowered_name:
        return "resume"

    lowered_text = text_content.lower()
    jd_signals = (
        "岗位职责",
        "任职要求",
        "岗位要求",
        "职位描述",
        "职位要求",
        "job description",
        "requirements",
        "responsibilities",
        "qualification",
    )
    resume_signals = (
        "工作经历",
        "项目经历",
        "实习经历",
        "教育背景",
        "教育经历",
        "个人总结",
        "自我评价",
        "技能",
        "experience",
        "education",
        "skills",
        "summary",
    )

    jd_score = sum(1 for signal in jd_signals if signal in lowered_text)
    resume_score = sum(1 for signal in resume_signals if signal in lowered_text)
    if jd_score > resume_score:
        return "jd"
    if resume_score > jd_score:
        return "resume"
    return "other"


def _persist_upload(*, raw: bytes, filename: str, session_id: str) -> Path:
    """Persist uploaded file under project artifacts or a shared upload folder."""
    safe_name = Path(filename).name or "upload.txt"
    if session_id:
        session = agent_runtime.session_crud.get(session_id)
        if session is None:
            raise HTTPException(status_code=404, detail=f"Session not found: {session_id}")
        target_dir = Path(".data") / "artifacts" / session.project_id / "uploads"
    else:
        target_dir = Path(".data") / "uploads"
    target_dir.mkdir(parents=True, exist_ok=True)
    target_path = target_dir / safe_name
    if target_path.exists():
        stem = Path(safe_name).stem
        suffix = Path(safe_name).suffix
        index = 1
    while target_path.exists():
        target_path = target_dir / f"{stem}_{index}{suffix}"
        index += 1
    target_path.write_bytes(raw)
    return target_path.resolve()


if __name__ == "__main__":
    import uvicorn

    uvicorn.run("api.main:app", host="127.0.0.1", port=8000, reload=False)
